import base64
import difflib
import io
import json
import os
import re
import tempfile
import threading
import time
import uuid
from typing import Any, Callable, Optional

import bleach
import markdown
import requests
from bs4 import BeautifulSoup
from docx import Document
from flask import Flask, jsonify, render_template, request, send_file
from openpyxl import Workbook
from openpyxl.styles import Font
from pypdf import PdfReader, PdfWriter


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 30 * 1024 * 1024  # 30 MB

TYPHOON_OCR_URL = "https://api.opentyphoon.ai/v1/ocr"
OCR_JOBS: dict[str, dict[str, Any]] = {}
OCR_JOBS_LOCK = threading.Lock()
ALLOWED_HTML_TAGS = set(bleach.sanitizer.ALLOWED_TAGS).union(
    {
        "p",
        "br",
        "pre",
        "code",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "table",
        "thead",
        "tbody",
        "tr",
        "th",
        "td",
    }
)


def render_ocr_html(text: str) -> str:
    html = markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
    )
    return bleach.clean(html, tags=ALLOWED_HTML_TAGS, strip=True)


def parse_tables_from_html(html: str) -> list[list[list[str]]]:
    soup = BeautifulSoup(html or "", "html.parser")
    parsed_tables = []
    for table in soup.find_all("table"):
        rows = []
        for tr in table.find_all("tr"):
            cells = tr.find_all(["th", "td"])
            if not cells:
                continue
            row = [" ".join(cell.stripped_strings) for cell in cells]
            rows.append(row)
        if rows:
            parsed_tables.append(rows)
    return parsed_tables


def normalize_header_text(value: str) -> str:
    normalized = re.sub(r"\s+", "", (value or "").lower())
    normalized = re.sub(r"[^0-9a-zก-๙]+", "", normalized)
    return normalized


def header_similarity(left: str, right: str) -> float:
    left_norm = normalize_header_text(left)
    right_norm = normalize_header_text(right)
    if not left_norm or not right_norm:
        return 0.0
    return difflib.SequenceMatcher(None, left_norm, right_norm).ratio()


def is_row_similar_to_header(row: list[str], header: list[str]) -> bool:
    if not row or not header:
        return False
    limit = min(len(row), len(header))
    if limit == 0:
        return False
    matches = 0
    for idx in range(limit):
        if header_similarity(row[idx], header[idx]) >= 0.8:
            matches += 1
    return (matches / limit) >= 0.6


def build_column_mapping(source_header: list[str], target_header: list[str]) -> dict[int, int]:
    mapping: dict[int, int] = {}
    used_targets: set[int] = set()

    for src_idx, src_name in enumerate(source_header):
        best_target = None
        best_score = 0.0
        for tgt_idx, tgt_name in enumerate(target_header):
            if tgt_idx in used_targets:
                continue
            score = header_similarity(src_name, tgt_name)
            if score > best_score:
                best_score = score
                best_target = tgt_idx

        if best_target is not None and best_score >= 0.45:
            mapping[src_idx] = best_target
            used_targets.add(best_target)
        elif src_idx < len(target_header):
            # Fallback by position if header OCR is too noisy.
            mapping[src_idx] = src_idx

    return mapping


def align_row_to_header(
    row: list[str],
    mapping: dict[int, int],
    target_width: int,
) -> list[str]:
    aligned = [""] * target_width
    for src_idx, value in enumerate(row):
        target_idx = mapping.get(src_idx)
        if target_idx is None or target_idx >= target_width:
            continue
        value_clean = (value or "").strip()
        if not value_clean:
            continue
        if aligned[target_idx]:
            aligned[target_idx] = f"{aligned[target_idx]} {value_clean}"
        else:
            aligned[target_idx] = value_clean
    return aligned


def merge_table_rows(tables: list[list[list[str]]]) -> list[list[str]]:
    if not tables:
        return []

    first_table = tables[0]
    if not first_table:
        return []

    primary_header = first_table[0]
    merged_rows: list[list[str]] = [primary_header]
    target_width = len(primary_header)

    first_mapping = {idx: idx for idx in range(target_width)}
    for row in first_table[1:]:
        if is_row_similar_to_header(row, primary_header):
            continue
        merged_rows.append(align_row_to_header(row, first_mapping, target_width))

    for table_rows in tables[1:]:
        if not table_rows:
            continue
        source_header = table_rows[0]
        mapping = build_column_mapping(source_header, primary_header)
        body_rows = table_rows[1:]

        for row in body_rows:
            if is_row_similar_to_header(row, source_header) or is_row_similar_to_header(row, primary_header):
                continue
            merged_rows.append(align_row_to_header(row, mapping, target_width))

    return merged_rows


def build_source_table_payloads(page_htmls: list[str], fallback_html: str) -> list[dict]:
    payloads: list[dict] = []

    if page_htmls:
        for page_number, page_html in enumerate(page_htmls, start=1):
            for table_rows in parse_tables_from_html(page_html):
                if not table_rows:
                    continue
                header = table_rows[0]
                row_order = 0
                rows_with_source = []
                for row in table_rows[1:]:
                    if is_row_similar_to_header(row, header):
                        continue
                    row_order += 1
                    rows_with_source.append((row, f"{page_number}-{row_order}"))
                payloads.append({"header": header, "rows": rows_with_source})
        return payloads

    # Fallback when per-page html is unavailable.
    for table_rows in parse_tables_from_html(fallback_html):
        if not table_rows:
            continue
        header = table_rows[0]
        row_order = 0
        rows_with_source = []
        for row in table_rows[1:]:
            if is_row_similar_to_header(row, header):
                continue
            row_order += 1
            rows_with_source.append((row, f"1-{row_order}"))
        payloads.append({"header": header, "rows": rows_with_source})

    return payloads


def merge_table_rows_with_source(table_payloads: list[dict]) -> list[list[str]]:
    if not table_payloads:
        return []

    primary_header = table_payloads[0]["header"]
    target_width = len(primary_header)
    merged_rows: list[list[str]] = [primary_header + ["ที่มาของข้อมูล"]]

    for index, payload in enumerate(table_payloads):
        source_header = payload.get("header", [])
        rows_with_source = payload.get("rows", [])
        if not source_header:
            continue

        if index == 0:
            mapping = {idx: idx for idx in range(target_width)}
        else:
            mapping = build_column_mapping(source_header, primary_header)

        for row, source_ref in rows_with_source:
            if is_row_similar_to_header(row, source_header) or is_row_similar_to_header(row, primary_header):
                continue
            aligned = align_row_to_header(row, mapping, target_width)
            merged_rows.append(aligned + [source_ref])

    return merged_rows


def decode_base64_payload(value: str) -> str:
    if not value:
        return ""
    try:
        return base64.b64decode(value.encode("utf-8")).decode("utf-8")
    except Exception:
        return ""


def decode_base64_json_list(value: str) -> list[str]:
    decoded = decode_base64_payload(value)
    if not decoded:
        return []
    try:
        parsed = json.loads(decoded)
        if isinstance(parsed, list):
            return [str(item) for item in parsed]
    except Exception:
        pass
    return []


def parse_pages_input(pages_raw: str) -> Optional[list[int]]:
    """
    Parse page input to Typhoon pages JSON.
    Supports: "", "all", "1,2,3", "1-5", "1-3,8,10-12"
    Returns list of page numbers or None for all pages.
    """
    raw = (pages_raw or "").strip().lower()
    if not raw or raw == "all":
        return None

    normalized = raw.replace(" ", "")
    tokens = [token for token in normalized.split(",") if token]
    if not tokens:
        return None

    pages_set = set()
    for token in tokens:
        if "-" in token:
            parts = token.split("-")
            if len(parts) != 2 or not parts[0].isdigit() or not parts[1].isdigit():
                raise ValueError("รูปแบบ Pages ไม่ถูกต้อง (ตัวอย่างที่ถูก: 1-39 หรือ 1,2,3)")
            start = int(parts[0])
            end = int(parts[1])
            if start <= 0 or end <= 0:
                raise ValueError("เลขหน้าใน Pages ต้องมากกว่า 0")
            if start > end:
                raise ValueError("ช่วงหน้าใน Pages ต้องเรียงจากน้อยไปมาก เช่น 1-39")
            for page in range(start, end + 1):
                pages_set.add(page)
        else:
            if not token.isdigit():
                raise ValueError("รูปแบบ Pages ไม่ถูกต้อง (ตัวอย่างที่ถูก: 1-39 หรือ 1,2,3)")
            page = int(token)
            if page <= 0:
                raise ValueError("เลขหน้าใน Pages ต้องมากกว่า 0")
            pages_set.add(page)

    if not pages_set:
        return None

    return sorted(pages_set)


def get_pdf_page_count(file_path: str) -> int:
    reader = PdfReader(file_path)
    return len(reader.pages)


def _parse_native_table_rows(lines: list[str]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    current_vin = ""
    date_range_pattern = r"\d{1,2}/\d{1,2}/\d{4}-\d{1,2}/\d{1,2}/\d{4}"
    money_pattern = r"[0-9,]+\.[0-9]{2}"
    rate_pattern = r"[0-9]+\.[0-9]+"

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        vin_match = re.match(r"^([A-Z0-9]{10,})\s*\(Closed", line, flags=re.IGNORECASE)
        if vin_match:
            current_vin = vin_match.group(1)
            continue

        if not line.startswith("ดอกเบี้ย"):
            continue

        row_pattern = (
            rf"^ดอกเบี้ย\s+({date_range_pattern})\s+(\d+)\s+({rate_pattern})\s+({money_pattern})\s+({money_pattern})$"
        )
        row_match = re.match(row_pattern, line)
        if not row_match:
            continue

        if not current_vin:
            continue

        rows.append(
            {
                "vin": current_vin,
                "item": "ดอกเบี้ย",
                "period": row_match.group(1),
                "days": row_match.group(2),
                "rate": row_match.group(3),
                "principal": row_match.group(4),
                "amount": row_match.group(5),
            }
        )

    return rows


def extract_native_page_content(file_path: str, page_number: int) -> tuple[str, int]:
    reader = PdfReader(file_path)
    if page_number < 1 or page_number > len(reader.pages):
        return "", 0

    text = reader.pages[page_number - 1].extract_text() or ""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "", 0

    rows = _parse_native_table_rows(lines)
    if not rows:
        # No parseable table, but keep native text as fallback.
        return "\n".join(lines), 0

    header_lines: list[str] = []
    for line in lines:
        if "เลขตัวถัง/เลขทะเบียน" in line or line.startswith("PRB") or line.startswith("LNN") or line.startswith("LVU"):
            break
        if line in {"-- 1 of 1 --"}:
            continue
        header_lines.append(line)

    md_lines = []
    if header_lines:
        md_lines.append("\n".join(header_lines))
        md_lines.append("")

    md_lines.append("| เลขตัวถัง/เลขทะเบียน | รายการ | ระยะเวลา | วัน | อัตราดอกเบี้ย | ต้นเงินกู้/เงินต้นคงเหลือ | จำนวนเงินที่ต้องชำระ |")
    md_lines.append("|---|---|---|---:|---:|---:|---:|")
    for row in rows:
        md_lines.append(
            f"| {row['vin']} (Closed) | {row['item']} | {row['period']} | {row['days']} | {row['rate']} | {row['principal']} | {row['amount']} |"
        )

    return "\n".join(md_lines), len(rows)


def init_ocr_job(job_id: str) -> None:
    with OCR_JOBS_LOCK:
        OCR_JOBS[job_id] = {
            "status": "running",
            "message": "กำลังเตรียมไฟล์",
            "current_step": 0,
            "total_steps": 0,
            "current_page_number": 0,
            "page_timings": [],
            "error": "",
            "result": None,
        }


def update_ocr_job(
    job_id: str,
    *,
    status: Optional[str] = None,
    message: Optional[str] = None,
    current_step: Optional[int] = None,
    total_steps: Optional[int] = None,
    current_page_number: Optional[int] = None,
    error: Optional[str] = None,
    result: Optional[dict] = None,
) -> None:
    with OCR_JOBS_LOCK:
        job = OCR_JOBS.get(job_id)
        if not job:
            return
        if status is not None:
            job["status"] = status
        if message is not None:
            job["message"] = message
        if current_step is not None:
            job["current_step"] = current_step
        if total_steps is not None:
            job["total_steps"] = total_steps
        if current_page_number is not None:
            job["current_page_number"] = current_page_number
        if error is not None:
            job["error"] = error
        if result is not None:
            job["result"] = result


def append_ocr_job_page_timing(job_id: str, page_number: int, elapsed_seconds: float) -> None:
    with OCR_JOBS_LOCK:
        job = OCR_JOBS.get(job_id)
        if not job:
            return
        timings = job.setdefault("page_timings", [])
        timings.append(
            {
                "page_number": page_number,
                "elapsed_seconds": round(float(elapsed_seconds), 2),
            }
        )


def export_tables_to_docx(
    html: str,
    fallback_text: str,
    page_htmls: Optional[list[str]] = None,
) -> io.BytesIO:
    table_payloads = build_source_table_payloads(page_htmls or [], html)
    merged_rows = merge_table_rows_with_source(table_payloads)

    doc = Document()
    doc.add_heading("OCR Result", level=1)

    if merged_rows:
        max_cols = max(len(row) for row in merged_rows)
        table = doc.add_table(rows=len(merged_rows), cols=max_cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(merged_rows):
            for c_idx in range(max_cols):
                value = row[c_idx] if c_idx < len(row) else ""
                cell = table.cell(r_idx, c_idx)
                cell.text = value
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
    else:
        doc.add_paragraph("No table found in OCR result.")
        if fallback_text.strip():
            doc.add_paragraph(fallback_text)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def export_tables_to_excel(
    html: str,
    fallback_text: str,
    page_htmls: Optional[list[str]] = None,
) -> io.BytesIO:
    table_payloads = build_source_table_payloads(page_htmls or [], html)
    merged_rows = merge_table_rows_with_source(table_payloads)

    wb = Workbook()
    ws = wb.active
    ws.title = "OCR Tables"

    if merged_rows:
        max_cols = max(len(row) for row in merged_rows)
        for row in merged_rows:
            padded = row + [""] * (max_cols - len(row))
            ws.append(padded)
        for cell in ws[1]:
            cell.font = Font(bold=True)
    else:
        ws.append(["No table found in OCR result."])
        if fallback_text.strip():
            ws.append([fallback_text])

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def save_unlocked_pdf(uploaded_bytes: bytes, password: Optional[str]) -> str:
    """Return path to a temporary unlocked PDF file."""
    reader = PdfReader(io.BytesIO(uploaded_bytes))

    if reader.is_encrypted:
        if not password:
            raise ValueError("ไฟล์ PDF นี้ถูกล็อก กรุณาใส่รหัสผ่าน PDF")
        if reader.decrypt(password) == 0:
            raise ValueError("รหัสผ่าน PDF ไม่ถูกต้อง")

    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        writer.write(tmp_file)
        return tmp_file.name


def call_typhoon_ocr_single_request(
    file_path: str,
    api_key: str,
    model: str,
    task_type: str,
    max_tokens: int,
    temperature: float,
    top_p: float,
    repetition_penalty: float,
    pages_json: Optional[str] = None,
) -> tuple[str, list[str]]:
    data = {
        "model": model,
        "task_type": task_type,
        "max_tokens": str(max_tokens),
        "temperature": str(temperature),
        "top_p": str(top_p),
        "repetition_penalty": str(repetition_penalty),
    }
    if pages_json:
        data["pages"] = pages_json

    headers = {"Authorization": f"Bearer {api_key}"}

    with open(file_path, "rb") as file:
        response = requests.post(
            TYPHOON_OCR_URL,
            files={"file": file},
            data=data,
            headers=headers,
            timeout=180,
        )

    if response.status_code != 200:
        raise RuntimeError(f"Typhoon API error {response.status_code}: {response.text}")

    result = response.json()
    extracted_texts = []
    per_page_texts: list[str] = []
    for page_result in result.get("results", []):
        if page_result.get("success") and page_result.get("message"):
            content = page_result["message"]["choices"][0]["message"]["content"]
            try:
                parsed_content = json.loads(content)
                text = parsed_content.get("natural_text", content)
            except json.JSONDecodeError:
                text = content
            extracted_texts.append(text)
            per_page_texts.append(text)
        elif not page_result.get("success"):
            error_msg = page_result.get("error", "Unknown error")
            extracted_texts.append(f"[ERROR] {error_msg}")
            per_page_texts.append(f"[ERROR] {error_msg}")

    return "\n\n".join(extracted_texts), per_page_texts


def call_typhoon_ocr(
    file_path: str,
    api_key: str,
    model: str,
    task_type: str,
    max_tokens: int,
    temperature: float,
    top_p: float,
    repetition_penalty: float,
    pages: Optional[list[int]] = None,
    progress_callback: Optional[Callable[[int, int, int], None]] = None,
    page_done_callback: Optional[Callable[[int, int, int, float], None]] = None,
) -> tuple[str, list[str], list[dict[str, Any]]]:
    """
    OCR strategy:
    - If multiple pages are requested, call API page-by-page for better completeness.
    - If one page (or unknown page list), do a single call.
    """
    if pages and len(pages) > 1:
        merged_texts: list[str] = []
        per_page_texts: list[str] = []
        page_timings: list[dict[str, Any]] = []
        total_pages = len(pages)

        for page_index, page_number in enumerate(pages, start=1):
            if progress_callback:
                progress_callback(page_index, total_pages, page_number)
            page_started = time.perf_counter()

            native_text, native_row_count = extract_native_page_content(file_path, page_number)
            if native_text and native_row_count > 0:
                merged_texts.append(native_text)
                per_page_texts.append(native_text)
                elapsed = round(time.perf_counter() - page_started, 2)
                page_timings.append({"page_number": page_number, "elapsed_seconds": elapsed})
                if page_done_callback:
                    page_done_callback(page_index, total_pages, page_number, elapsed)
                continue

            page_payload = json.dumps([page_number])
            last_error: Optional[Exception] = None

            # Retry each page a few times to reduce transient misses/timeouts.
            for _ in range(3):
                try:
                    page_joined, page_results = call_typhoon_ocr_single_request(
                        file_path=file_path,
                        api_key=api_key,
                        model=model,
                        task_type=task_type,
                        max_tokens=max_tokens,
                        temperature=temperature,
                        top_p=top_p,
                        repetition_penalty=repetition_penalty,
                        pages_json=page_payload,
                    )
                    if page_results:
                        page_text = page_results[0]
                    else:
                        page_text = page_joined
                    merged_texts.append(page_text)
                    per_page_texts.append(page_text)
                    last_error = None
                    break
                except Exception as exc:
                    last_error = exc

            if last_error is not None:
                raise RuntimeError(f"OCR failed on page {page_number}: {last_error}") from last_error

            elapsed = round(time.perf_counter() - page_started, 2)
            page_timings.append({"page_number": page_number, "elapsed_seconds": elapsed})
            if page_done_callback:
                page_done_callback(page_index, total_pages, page_number, elapsed)

        return "\n\n".join(merged_texts), per_page_texts, page_timings

    if progress_callback:
        single_page_number = pages[0] if pages else 1
        progress_callback(1, 1, single_page_number)

    started = time.perf_counter()
    if pages and len(pages) == 1:
        native_text, native_row_count = extract_native_page_content(file_path, pages[0])
        if native_text and native_row_count > 0:
            elapsed = round(time.perf_counter() - started, 2)
            page_timings = [{"page_number": pages[0], "elapsed_seconds": elapsed}]
            if page_done_callback:
                page_done_callback(1, 1, pages[0], elapsed)
            return native_text, [native_text], page_timings

    pages_json = json.dumps(pages) if pages else None
    merged_text, per_page_texts = call_typhoon_ocr_single_request(
        file_path=file_path,
        api_key=api_key,
        model=model,
        task_type=task_type,
        max_tokens=max_tokens,
        temperature=temperature,
        top_p=top_p,
        repetition_penalty=repetition_penalty,
        pages_json=pages_json,
    )

    elapsed = round(time.perf_counter() - started, 2)
    page_number = pages[0] if pages else 1
    page_timings = [{"page_number": page_number, "elapsed_seconds": elapsed}]
    if page_done_callback:
        page_done_callback(1, 1, page_number, elapsed)
    return merged_text, per_page_texts, page_timings


def run_ocr_pipeline(
    *,
    uploaded_bytes: bytes,
    pdf_password: str,
    api_key: str,
    model: str,
    task_type: str,
    max_tokens: int,
    temperature: float,
    top_p: float,
    repetition_penalty: float,
    pages_raw: str,
    progress_callback: Optional[Callable[[int, int, int], None]] = None,
    page_done_callback: Optional[Callable[[int, int, int, float], None]] = None,
) -> dict[str, Any]:
    temp_pdf_path = ""
    start_time = time.perf_counter()
    try:
        temp_pdf_path = save_unlocked_pdf(uploaded_bytes, pdf_password)

        if pages_raw:
            pages_value = parse_pages_input(pages_raw)
            if pages_value is None:
                page_count = get_pdf_page_count(temp_pdf_path)
                pages_value = list(range(1, page_count + 1))
        else:
            page_count = get_pdf_page_count(temp_pdf_path)
            pages_value = list(range(1, page_count + 1))

        extracted_text, page_texts, page_timings = call_typhoon_ocr(
            file_path=temp_pdf_path,
            api_key=api_key,
            model=model,
            task_type=task_type,
            max_tokens=max_tokens,
            temperature=temperature,
            top_p=top_p,
            repetition_penalty=repetition_penalty,
            pages=pages_value,
            progress_callback=progress_callback,
            page_done_callback=page_done_callback,
        )

        extracted_html = render_ocr_html(extracted_text)
        page_htmls = [render_ocr_html(page_text) for page_text in page_texts]
        extracted_text_b64 = base64.b64encode(extracted_text.encode("utf-8")).decode("utf-8")
        extracted_html_b64 = base64.b64encode(extracted_html.encode("utf-8")).decode("utf-8")
        page_htmls_b64 = base64.b64encode(
            json.dumps(page_htmls, ensure_ascii=False).encode("utf-8")
        ).decode("utf-8")
        elapsed_seconds = round(time.perf_counter() - start_time, 2)

        return {
            "extracted_text": extracted_text,
            "extracted_html": extracted_html,
            "extracted_text_b64": extracted_text_b64,
            "extracted_html_b64": extracted_html_b64,
            "page_htmls_b64": page_htmls_b64,
            "page_texts": page_texts,
            "page_htmls": page_htmls,
            "page_timings": page_timings,
            "elapsed_seconds": elapsed_seconds,
        }
    finally:
        if temp_pdf_path and os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)


def run_ocr_job(job_id: str, params: dict[str, Any]) -> None:
    try:
        def on_progress(current_step: int, total_steps: int, page_number: int) -> None:
            update_ocr_job(
                job_id,
                status="running",
                message=f"กำลัง OCR หน้า {current_step}/{total_steps}",
                current_step=current_step,
                total_steps=total_steps,
                current_page_number=page_number,
            )

        def on_page_done(
            current_step: int,
            total_steps: int,
            page_number: int,
            elapsed_seconds: float,
        ) -> None:
            append_ocr_job_page_timing(job_id, page_number, elapsed_seconds)
            update_ocr_job(
                job_id,
                status="running",
                message=f"เสร็จหน้า {current_step}/{total_steps} (หน้าเอกสาร {page_number})",
                current_step=current_step,
                total_steps=total_steps,
                current_page_number=page_number,
            )

        result = run_ocr_pipeline(
            uploaded_bytes=params["uploaded_bytes"],
            pdf_password=params["pdf_password"],
            api_key=params["api_key"],
            model=params["model"],
            task_type=params["task_type"],
            max_tokens=params["max_tokens"],
            temperature=params["temperature"],
            top_p=params["top_p"],
            repetition_penalty=params["repetition_penalty"],
            pages_raw=params["pages_raw"],
            progress_callback=on_progress,
            page_done_callback=on_page_done,
        )
        update_ocr_job(
            job_id,
            status="completed",
            message="OCR เสร็จแล้ว",
            result=result,
        )
    except Exception as exc:
        update_ocr_job(
            job_id,
            status="failed",
            message="OCR ล้มเหลว",
            error=str(exc),
        )


@app.route("/ocr/start", methods=["POST"])
def ocr_start():
    uploaded_file = request.files.get("pdf_file")
    if not uploaded_file or uploaded_file.filename == "":
        return jsonify({"ok": False, "error": "กรุณาเลือกไฟล์ PDF ก่อน"}), 400

    api_key = request.form.get("api_key", "").strip() or os.getenv("TYPHOON_API_KEY", "")
    if not api_key:
        return jsonify({"ok": False, "error": "กรุณาใส่ Typhoon API Key"}), 400

    params = {
        "uploaded_bytes": uploaded_file.read(),
        "pdf_password": request.form.get("pdf_password", "").strip(),
        "api_key": api_key,
        "model": request.form.get("model", "typhoon-ocr").strip(),
        "task_type": request.form.get("task_type", "default").strip(),
        "max_tokens": int(request.form.get("max_tokens", "16384")),
        "temperature": float(request.form.get("temperature", "0.1")),
        "top_p": float(request.form.get("top_p", "0.6")),
        "repetition_penalty": float(request.form.get("repetition_penalty", "1.2")),
        "pages_raw": request.form.get("pages", "").strip(),
    }

    job_id = uuid.uuid4().hex
    init_ocr_job(job_id)
    worker = threading.Thread(target=run_ocr_job, args=(job_id, params), daemon=True)
    worker.start()
    return jsonify({"ok": True, "job_id": job_id})


@app.route("/ocr/status/<job_id>", methods=["GET"])
def ocr_status(job_id: str):
    with OCR_JOBS_LOCK:
        job = OCR_JOBS.get(job_id)
        if not job:
            return jsonify({"ok": False, "error": "ไม่พบงาน OCR"}), 404
        response = {
            "ok": True,
            "status": job["status"],
            "message": job.get("message", ""),
            "current_step": job.get("current_step", 0),
            "total_steps": job.get("total_steps", 0),
            "current_page_number": job.get("current_page_number", 0),
            "page_timings": job.get("page_timings", []),
            "error": job.get("error", ""),
        }
        if job["status"] == "completed" and job.get("result") is not None:
            response["result"] = job["result"]
    return jsonify(response)


@app.route("/", methods=["GET", "POST"])
def index():
    extracted_text = ""
    extracted_html = ""
    extracted_text_b64 = ""
    extracted_html_b64 = ""
    page_htmls_b64 = ""
    page_texts = []
    page_htmls = []
    page_timings = []
    elapsed_seconds = None
    error = ""

    defaults = {
        "model": "typhoon-ocr",
        "task_type": "default",
        "max_tokens": "16384",
        "temperature": "0.1",
        "top_p": "0.6",
        "repetition_penalty": "1.2",
    }

    if request.method == "POST":
        uploaded_file = request.files.get("pdf_file")
        pdf_password = request.form.get("pdf_password", "").strip()
        api_key = request.form.get("api_key", "").strip() or os.getenv("TYPHOON_API_KEY", "")

        model = request.form.get("model", defaults["model"]).strip()
        task_type = request.form.get("task_type", defaults["task_type"]).strip()
        max_tokens = int(request.form.get("max_tokens", defaults["max_tokens"]))
        temperature = float(request.form.get("temperature", defaults["temperature"]))
        top_p = float(request.form.get("top_p", defaults["top_p"]))
        repetition_penalty = float(
            request.form.get("repetition_penalty", defaults["repetition_penalty"])
        )
        pages_raw = request.form.get("pages", "").strip()

        if not uploaded_file or uploaded_file.filename == "":
            error = "กรุณาเลือกไฟล์ PDF ก่อน"
        elif not api_key:
            error = "กรุณาใส่ Typhoon API Key"
        else:
            try:
                result = run_ocr_pipeline(
                    uploaded_bytes=uploaded_file.read(),
                    pdf_password=pdf_password,
                    api_key=api_key,
                    model=model,
                    task_type=task_type,
                    max_tokens=max_tokens,
                    temperature=temperature,
                    top_p=top_p,
                    repetition_penalty=repetition_penalty,
                    pages_raw=pages_raw,
                )
                extracted_text = result["extracted_text"]
                extracted_html = result["extracted_html"]
                extracted_text_b64 = result["extracted_text_b64"]
                extracted_html_b64 = result["extracted_html_b64"]
                page_htmls_b64 = result["page_htmls_b64"]
                page_texts = result["page_texts"]
                page_htmls = result["page_htmls"]
                page_timings = result["page_timings"]
                elapsed_seconds = result["elapsed_seconds"]
            except Exception as exc:  # keep UI simple
                error = str(exc)

    return render_template(
        "index.html",
        extracted_text=extracted_text,
        extracted_html=extracted_html,
        extracted_text_b64=extracted_text_b64,
        extracted_html_b64=extracted_html_b64,
        page_htmls_b64=page_htmls_b64,
        page_texts=page_texts,
        page_htmls=page_htmls,
        page_timings=page_timings,
        elapsed_seconds=elapsed_seconds,
        error=error,
        defaults=defaults,
    )


@app.route("/download/word", methods=["POST"])
def download_word():
    extracted_html = decode_base64_payload(request.form.get("extracted_html_b64", ""))
    extracted_text = decode_base64_payload(request.form.get("extracted_text_b64", ""))
    page_htmls = decode_base64_json_list(request.form.get("page_htmls_b64", ""))
    file_data = export_tables_to_docx(extracted_html, extracted_text, page_htmls=page_htmls)
    return send_file(
        file_data,
        as_attachment=True,
        download_name="ocr-tables.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/download/excel", methods=["POST"])
def download_excel():
    extracted_html = decode_base64_payload(request.form.get("extracted_html_b64", ""))
    extracted_text = decode_base64_payload(request.form.get("extracted_text_b64", ""))
    page_htmls = decode_base64_json_list(request.form.get("page_htmls_b64", ""))
    file_data = export_tables_to_excel(extracted_html, extracted_text, page_htmls=page_htmls)
    return send_file(
        file_data,
        as_attachment=True,
        download_name="ocr-tables.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)
